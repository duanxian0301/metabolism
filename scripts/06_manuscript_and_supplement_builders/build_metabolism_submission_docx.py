from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import unescape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"D:\codex\GenomicSEM\metabolic")
MANUSCRIPT = ROOT / "manuscript" / "metabolic_factor_ndd_main_manuscript_integrated_v3_S24b.md"
BIB = ROOT / "manuscript" / "references.bib"
OUT_MD = ROOT / "manuscript" / "metabolic_factor_ndd_metabolism_submission_v2.md"
OUT_DOCX = ROOT / "manuscript" / "metabolic_factor_ndd_metabolism_submission_v2.docx"
FIG_DIR = ROOT / "figures"


TITLE = (
    "Metabolic Genetic Axes Linked to Alzheimer and Parkinson Disease Through "
    "Distinct Polygenic, Cellular, and Target-Annotation Programs"
)

AUTHORS = (
    "Jing Shen1†, Yuxuan Shen1†, Gang Yu1†, Lijun Pang1*, Chenxu Xiao1,*"
)

AFFILIATIONS = [
    "1 Central laboratory, The Affiliated Jiangsu Shengze Hospital of Nanjing Medical University, Suzhou, Jiangsu, China.",
    "†Co-first author: Jing Shen, Email: shenjing@njmu.edu.cn; ORCID ID: 0009-0001-5120-9897.",
    "†Co-first author: Yuxuan Shen, Email: 353160163@qq.com.",
    "†Co-first author: Gang Yu, Email: 550545669@qq.com.",
    "*Corresponding authors: Lijun Pang, Email: pplljj1986@126.com; Chenxu Xiao, Email: xcx5443@gmail.com.",
]

ABSTRACT = {
    "Background": (
        "Metabolic dysregulation is implicated in Alzheimer disease (AD) and Parkinson disease (PD), "
        "but individual circulating metabolites are highly correlated and may not capture shared metabolic programs."
    ),
    "Methods": (
        "We applied Genomic SEM to NMR metabolite GWAS summary statistics to construct lipid and non-lipid "
        "latent metabolic factor GWAS. Six factors were screened against AD, PD, and Lewy body dementia using "
        "cross-trait LDSC, followed by MiXeR, pleiotropy-informed locus discovery, colocalization, transcriptomic "
        "prioritization, single-cell projection, virtual knockout, and target annotation."
    ),
    "Results": (
        "Among 18 factor-disease comparisons, three survived FDR correction: an HDL-core axis positively correlated "
        "with AD (rg = 0.142, FDR = 0.00234), a ketone-body axis inversely correlated with PD (rg = -0.106, "
        "FDR = 0.0154), and a TG-rich/VLDL-remodeling axis inversely correlated with PD (rg = -0.0719, "
        "FDR = 0.0308). Component-level LDSC supported the biological traceability of these axes. Downstream analyses "
        "linked HDL-core/AD to perivascular support, ketone-body/PD to distributed neuronal-glial energy adaptation, "
        "and TG/VLDL-remodeling/PD to lysosomal and membrane-trafficking programs. Target annotation prioritized "
        "experimentally approachable genes including GRK4, PRMT7, and TMEM175."
    ),
    "Conclusions": (
        "Structured metabolic genetic axes reveal selective and biologically distinct links between circulating "
        "metabolism and neurodegenerative disease, providing a framework for mechanism-focused follow-up beyond "
        "single-metabolite association."
    ),
}

KEYWORDS = [
    "Genomic SEM",
    "metabolomics",
    "Alzheimer disease",
    "Parkinson disease",
    "lipid metabolism",
    "ketone bodies",
    "single-cell genomics",
]

HIGHLIGHTS = [
    "Genomic SEM organized circulating metabolites into six genetic metabolic axes.",
    "HDL-core, ketone-body, and TG/VLDL axes aligned selectively with AD or PD.",
    "Single-cell analyses separated vascular, energy, and lysosomal programs.",
    "Target annotation prioritized GRK4, PRMT7, and TMEM175 for follow-up.",
]

ABBREVIATIONS = [
    ("AD", "Alzheimer disease"),
    ("APOE", "apolipoprotein E"),
    ("CFI", "comparative fit index"),
    ("conjFDR", "conjunctional false discovery rate"),
    ("cTWAS", "causal transcriptome-wide association study"),
    ("DGIdb", "Drug-Gene Interaction Database"),
    ("FDR", "false discovery rate"),
    ("FUMA", "Functional Mapping and Annotation"),
    ("GTEx", "Genotype-Tissue Expression"),
    ("GWAS", "genome-wide association study"),
    ("HDL", "high-density lipoprotein"),
    ("HPA", "Human Protein Atlas"),
    ("LBD", "Lewy body dementia"),
    ("LDSC", "linkage disequilibrium score regression"),
    ("MiXeR", "causal mixture model"),
    ("NDD", "neurodegenerative disease"),
    ("NMR", "nuclear magnetic resonance"),
    ("OPC", "oligodendrocyte precursor cell"),
    ("PD", "Parkinson disease"),
    ("PWCoCo", "pairwise colocalization"),
    ("Q_SNP", "SNP-level heterogeneity statistic"),
    ("rg", "genetic correlation"),
    ("scPagwas", "single-cell pathway-based genome-wide association study"),
    ("scTenifoldKnk", "single-cell TenifoldKnk"),
    ("SEM", "structural equation modelling"),
    ("SMR", "summary-data-based Mendelian randomization"),
    ("SNP", "single nucleotide polymorphism"),
    ("SRMR", "standardized root mean square residual"),
    ("TG", "triglyceride"),
    ("TRS", "trait-relevance score"),
    ("UMAP", "uniform manifold approximation and projection"),
    ("VLDL", "very-low-density lipoprotein"),
]

DECLARATIONS = {
    "Ethics approval and consent to participate": (
        "This study used publicly available or controlled-access summary-level genetic, transcriptomic, proteomic, "
        "and annotation resources. No new individual-level human participant data were collected. Ethical approval "
        "and informed consent were obtained by the original contributing studies."
    ),
    "Consent for publication": "Not applicable.",
    "Availability of data and materials": (
        "All source summary statistics and public resources used in this study are available from the original "
        "studies or databases cited in the manuscript. Derived analysis summaries are provided in the Supplementary "
        "Tables. Access-controlled source datasets remain subject to the policies of the original data providers."
    ),
    "Code availability": (
        "Analysis scripts used to generate the reported results are available from the corresponding authors upon "
        "reasonable request."
    ),
    "Competing interests": "The authors declare that they have no competing interests.",
    "Funding": "The authors declare that no specific funding was received for this work.",
    "Authors' contributions": (
        "Jing Shen: Conceptualization, Methodology, Formal analysis, Writing - original draft, Visualization. "
        "Yuxuan Shen: Data curation, Formal analysis, Validation, Visualization. "
        "Gang Yu: Software, Data curation, Validation, Investigation. "
        "Lijun Pang: Supervision, Project administration, Writing - review and editing. "
        "Chenxu Xiao: Conceptualization, Supervision, Methodology, Writing - review and editing, Project administration. "
        "All authors read and approved the final manuscript."
    ),
    "Acknowledgements": (
        "The authors thank the investigators and participants of the metabolomics, neurodegenerative disease GWAS, "
        "single-cell, eQTL, proteomic, transcriptomic, and target-annotation resources used in this study."
    ),
    "Declaration of generative AI and AI-assisted technologies in the writing process": (
        "During preparation of this work, the authors used AI-assisted tools to support language editing, formatting, "
        "reference organization, and document preparation. After using these tools, the authors reviewed and edited "
        "the content and take full responsibility for the scientific accuracy of the manuscript."
    ),
}

FIGURES = {
    2: {
        "path": FIG_DIR / "figure2_factor_models_disease_screen_v13_final_polished.png",
        "title": "Genetic construction, validation, and disease screening of lipid and non-lipid metabolic factors.",
        "legend": (
            "(A, B) GenomicSEM factor structures for the lipid and non-lipid metabolic domains. "
            "Metabolic traits are shown as observed indicators and latent metabolic factors as circles. "
            "Standardized factor loadings are shown on indicator paths, and latent factor correlations are shown between factors. "
            "(C) Genetic validation of the derived factor GWASs against their internal and supporting metabolic traits, summarized by bivariate LDSC genetic correlations. "
            "(D) Cross-trait LDSC screening between the six metabolic factor GWASs and neurodegenerative disease GWASs. "
            "Circle size represents the absolute genetic correlation (|rg|), and color represents the FDR-adjusted significance level. "
            "Three factor-disease pairs passed the FDR threshold and were taken forward: HDL-core with AD, ketone-core with PD, and TG/VLDL-core with PD. "
            "(E) MiXeR-based genetic-overlap estimates for the prioritized metabolic factor-disease pairs, showing the extent of shared polygenic architecture beyond genome-wide significant loci. "
            "AD, Alzheimer's disease; PD, Parkinson's disease; LDSC, linkage disequilibrium score regression; rg, genetic correlation; FDR, false discovery rate; MiXeR, causal mixture model."
        ),
    },
    3: {
        "path": FIG_DIR / "figure3_shared_loci_genes_v6_final_polished.png",
        "title": "Shared loci and convergent gene prioritization across the three metabolic factor-NDD axes.",
        "legend": (
            "(A) Summary of shared-locus discovery across the prioritized factor-disease pairs, including pleiotropic loci and downstream mapped genes. "
            "(B) Chromosomal distribution of shared loci identified for HDL-core/AD, ketone-core/PD, and TG/VLDL-core/PD. "
            "(C) Locus-level evidence supporting shared genetic signals, integrating conjunctional FDR, colocalization, PWCoCo, and FUMA annotation. "
            "Gene names in this panel denote nearest or mapped genes used to label shared regions rather than independent proof of gene causality. "
            "(D) Prioritized genes linked to shared-signal regions after integrating locus annotation with regulatory evidence, including FUMA mapping, SMR, and cTWAS. "
            "Colocalization and PWCoCo evidence in this panel indicate that the gene lies within or maps to a shared-signal region supported by locus-level evidence. "
            "AD, Alzheimer's disease; PD, Parkinson's disease; conjFDR, conjunctional false discovery rate; FUMA, Functional Mapping and Annotation; SMR, summary-data-based Mendelian randomization; cTWAS, causal transcriptome-wide association study; PWCoCo, pairwise colocalization."
        ),
    },
    4: {
        "path": FIG_DIR / "figure4_singlecell_pathway_v6_clean_labels.png",
        "title": "Single-cell localization and pathway-level convergence of prioritized metabolic factor-NDD signals.",
        "legend": (
            "(A) Regulatory evidence supporting prioritized genes across the three metabolic factor-disease axes, summarized by the number of genes supported by GTEx SMR, bulk-brain SMR, cell-type SMR, and cTWAS. "
            "(B) UMAP representation of the MSSM single-cell reference atlas used for scPagwas analyses, with major cell types annotated. "
            "(C-E) Single-cell TRS score UMAPs for HDL-core/AD, ketone-core/PD, and TG/VLDL-core/PD, respectively. "
            "Cells are colored from low to high TRS score, highlighting the cellular localization of genetic-metabolic signal enrichment. "
            "(F-H) Distribution of TRS scores across selected cell types for each prioritized factor-disease pair. Asterisks indicate cell-type enrichment passing FDR correction in scPagwas. "
            "(I, J) Pathway-level convergence between scPagwas and KNK perturbation analyses for AD and PD axes, respectively. "
            "Bars show top-50 rank scores for overlapping pathways, with higher values indicating stronger ranking within each method. Colored dots indicate the metabolic factor-disease axis contributing to each pathway. "
            "AD, Alzheimer's disease; PD, Parkinson's disease; UMAP, uniform manifold approximation and projection; TRS, trait-relevance score; scPagwas, single-cell pathway-based GWAS analysis; KNK, knockoff-based network knockout/perturbation analysis; SMR, summary-data-based Mendelian randomization; cTWAS, causal transcriptome-wide association study; FDR, false discovery rate."
        ),
    },
    5: {
        "path": FIG_DIR / "figure5_translational_prioritization_v2_smr_union.png",
        "title": "Translational prioritization of metabolic factor-linked candidate genes.",
        "legend": (
            "(A) Evidence-convergence matrix for prioritized target genes. Columns summarize genetic evidence, regulatory evidence, cell-level support, external disease-state support, and translational annotation. "
            "The SMR column represents the union of GTEx SMR, bulk-brain SMR, and cell-type SMR support. Filled squares indicate the presence of support from the corresponding evidence layer. "
            "(B) Supportive public disease-state summaries for prioritized genes. Circles show PD transcriptomic meta-signature support using the D.overall score, with darker purple indicating midbrain expression. "
            "Diamonds indicate genes present in the AD CSF proteomics summary. These analyses were used as supportive validation rather than primary causal evidence. "
            "(C) Target tractability and modality annotation for prioritized genes, integrating approved drug-gene evidence, small-molecule binding or pocket support, chemical probe availability, surface or antibody modality support, protein interaction evidence, and structural availability. "
            "Bar length represents the binding/modality support score, and color indicates the tractability tier. "
            "(D) Network-style summary linking the three prioritized metabolic factor-disease axes to candidate target genes. Node color reflects the metabolic axis or tractability tier, and edges indicate that a gene was prioritized within the corresponding axis. "
            "AD, Alzheimer's disease; PD, Parkinson's disease; SMR, summary-data-based Mendelian randomization; GTEx, Genotype-Tissue Expression project; CSF, cerebrospinal fluid."
        ),
    },
}


def parse_bibtex(path: Path) -> dict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    for m in re.finditer(r"@(\w+)\{([^,]+),", text):
        start = m.start()
        key = m.group(2)
        depth = 0
        end = None
        for i in range(start, len(text)):
            if text[i] == "{":
                depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    end = i + 1
                    break
        if end is None:
            continue
        body = text[m.end() : end - 1]
        fields = {"entry_type": m.group(1), "key": key}
        for fm in re.finditer(r"(\w+)\s*=\s*\{((?:[^{}]|\{[^{}]*\})*)\}\s*,?", body, flags=re.S):
            val = fm.group(2).replace("\n", " ")
            val = re.sub(r"\s+", " ", val).strip()
            val = val.replace("{", "").replace("}", "")
            fields[fm.group(1).lower()] = unescape(val)
        entries[key] = fields
    return entries


def extract_citation_order(text: str) -> list[str]:
    order: list[str] = []
    for cm in re.finditer(r"\[([^\]]*@[^]]+)\]", text):
        for key in re.findall(r"@([A-Za-z0-9_:-]+)", cm.group(1)):
            if key not in order:
                order.append(key)
    return order


def compress_numbers(nums: list[int]) -> str:
    nums = sorted(nums)
    ranges = []
    i = 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j > i:
            ranges.append(f"{nums[i]}-{nums[j]}")
        else:
            ranges.append(str(nums[i]))
        i = j + 1
    return ",".join(ranges)


def replace_citations(text: str, order: list[str]) -> str:
    idx = {k: i + 1 for i, k in enumerate(order)}

    def repl(m: re.Match[str]) -> str:
        keys = re.findall(r"@([A-Za-z0-9_:-]+)", m.group(1))
        nums = [idx[k] for k in keys if k in idx]
        return f"[{compress_numbers(nums)}]" if nums else m.group(0)

    return re.sub(r"\[([^\]]*@[^]]+)\]", repl, text)


def clean_md_body(text: str) -> str:
    text = text.replace("\ufeff", "")
    text = re.sub(r"^\s*# .+?\r?\n+", "", text)
    text = re.sub(r"\n## Reference Notes for Formatting[\s\S]*$", "\n", text)
    return text.strip()


def insert_figure_callouts(text: str) -> str:
    replacements = [
        (
            "Genome-wide factor and heterogeneity scans were summarized for all six factors. Factor-level genome-wide significant lead loci were observed across all six GWAS. Q_SNP results varied across factors, with no genome-wide heterogeneity for the HDL-core and cholesteryl-ester/structural lipid axes but appreciable heterogeneity for several remaining axes. We therefore treated Q_SNP as an annotation layer that distinguishes factor-consistent from potentially trait-specific loci (Supplementary Tables S19-S22).",
            "Genome-wide factor and heterogeneity scans were summarized for all six factors. Factor-level genome-wide significant lead loci were observed across all six GWAS. Q_SNP results varied across factors, with no genome-wide heterogeneity for the HDL-core and cholesteryl-ester/structural lipid axes but appreciable heterogeneity for several remaining axes. We therefore treated Q_SNP as an annotation layer that distinguishes factor-consistent from potentially trait-specific loci (Supplementary Tables S19-S22).\n\n[Insert Figure 2 near here]",
        ),
        (
            "Integrated gene prioritization reduced broad locus signals to candidate sets supported by convergent evidence. This was not a strict intersection across all methods; rather, genes were tiered by the breadth and consistency of support from positional/regulatory mapping, FUMA, cTWAS, bulk-brain SMR, GTEx brain-tissue SMR, and brain cell-type SMR. The HDL-core-AD branch retained 44 core genes, including YPEL3, INO80E, CAB39L, MED24, CDIPT, TBX6, LRRC37A2, ARL17B, DGKQ, HLA-DQB1, TMEM175, SLC26A1, and SLC12A9. The ketone-body-PD branch retained 14 core genes, with the highest-confidence subset including LRPAP1, DUS2, PRMT7, CENPV, GRK4, ZSWIM7, ELP5, and DDX55. The TG-rich/VLDL-remodeling-PD branch retained 50 core genes, including MAEA, PCGF3, SLC26A1, TMEM175, DGKQ, LRPAP1, LRRC37A2, CTSB, KAT8, GRK4, SHROOM3, ALKBH5, STX4, and WNT3 (Supplementary Tables S31-S41).",
            "Integrated gene prioritization reduced broad locus signals to candidate sets supported by convergent evidence. This was not a strict intersection across all methods; rather, genes were tiered by the breadth and consistency of support from positional/regulatory mapping, FUMA, cTWAS, bulk-brain SMR, GTEx brain-tissue SMR, and brain cell-type SMR. The HDL-core-AD branch retained 44 core genes, including YPEL3, INO80E, CAB39L, MED24, CDIPT, TBX6, LRRC37A2, ARL17B, DGKQ, HLA-DQB1, TMEM175, SLC26A1, and SLC12A9. The ketone-body-PD branch retained 14 core genes, with the highest-confidence subset including LRPAP1, DUS2, PRMT7, CENPV, GRK4, ZSWIM7, ELP5, and DDX55. The TG-rich/VLDL-remodeling-PD branch retained 50 core genes, including MAEA, PCGF3, SLC26A1, TMEM175, DGKQ, LRPAP1, LRRC37A2, CTSB, KAT8, GRK4, SHROOM3, ALKBH5, STX4, and WNT3 (Supplementary Tables S31-S41).\n\n[Insert Figure 3 near here]",
        ),
        (
            "Targeted scTenifoldKnk virtual knockout analyses further prioritized cell-context hypotheses. In HDL-core-AD, KANSL1 knockout in pericytes produced 73 perturbed genes and overlapped with cGMP-PKG, HIF-1, relaxin, transcriptional misregulation, and central carbon metabolism pathways. In ketone-body-PD, LRPAP1, PRMT7, GRK4, and DUS2 produced substantial perturbation profiles across pericytes, oligodendrocyte precursor cells, and VIP GABAergic interneurons, with recurrent overlap in synaptic, inflammatory, spliceosomal, and sodium-handling pathways. In TG-rich/VLDL-remodeling-PD, TMEM175 produced a compact pericyte perturbation profile linking ribosomal and platelet-activation pathways, whereas LRRC37A2 showed broader perturbation in oligodendrocyte precursor cells. These analyses were interpreted as prioritization of experimentally testable mechanisms (Supplementary Tables S48-S54).",
            "Targeted scTenifoldKnk virtual knockout analyses further prioritized cell-context hypotheses. In HDL-core-AD, KANSL1 knockout in pericytes produced 73 perturbed genes and overlapped with cGMP-PKG, HIF-1, relaxin, transcriptional misregulation, and central carbon metabolism pathways. In ketone-body-PD, LRPAP1, PRMT7, GRK4, and DUS2 produced substantial perturbation profiles across pericytes, oligodendrocyte precursor cells, and VIP GABAergic interneurons, with recurrent overlap in synaptic, inflammatory, spliceosomal, and sodium-handling pathways. In TG-rich/VLDL-remodeling-PD, TMEM175 produced a compact pericyte perturbation profile linking ribosomal and platelet-activation pathways, whereas LRRC37A2 showed broader perturbation in oligodendrocyte precursor cells. These analyses were interpreted as prioritization of experimentally testable mechanisms (Supplementary Tables S48-S54).\n\n[Insert Figure 4 near here]",
        ),
        (
            "Target annotation highlighted a focused set of experimentally approachable genes at the end of the discovery chain. GRK4 showed the strongest overall target profile, with approved drug-gene interaction evidence, kinase-family tractability, protein-interaction support, and available structure information. PRMT7 showed small-molecule and chemical-probe support, including SGC3027 and DS-437, and TMEM175 showed convergent membrane-target and small-molecule support. DUS2, DGKQ, LRPAP1, and LRRC37A2 showed intermediate tractability, whereas ARL17B, KANSL1, and KNOP1 were supported mainly by structure or network context with limited direct binding evidence (Supplementary Tables S55-S59). As a supportive disease-state check, all KNK-prioritized PD genes were present in a public PD midbrain-neuron transcriptomic meta-signature, with LRPAP1, GRK4, KNOP1, TMEM175, KANSL1, and PRMT7 among the higher-ranked prioritized genes (Supplementary Tables S60-S61; Supplementary Information).",
            "Target annotation highlighted a focused set of experimentally approachable genes at the end of the discovery chain. GRK4 showed the strongest overall target profile, with approved drug-gene interaction evidence, kinase-family tractability, protein-interaction support, and available structure information. PRMT7 showed small-molecule and chemical-probe support, including SGC3027 and DS-437, and TMEM175 showed convergent membrane-target and small-molecule support. DUS2, DGKQ, LRPAP1, and LRRC37A2 showed intermediate tractability, whereas ARL17B, KANSL1, and KNOP1 were supported mainly by structure or network context with limited direct binding evidence (Supplementary Tables S55-S59). As a supportive disease-state check, all KNK-prioritized PD genes were present in a public PD midbrain-neuron transcriptomic meta-signature, with LRPAP1, GRK4, KNOP1, TMEM175, KANSL1, and PRMT7 among the higher-ranked prioritized genes (Supplementary Tables S60-S61; Supplementary Information).\n\n[Insert Figure 5 near here]",
        ),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def format_authors(author_field: str | None) -> str:
    if not author_field:
        return ""
    author_field = author_field.replace("{", "").replace("}", "")
    parts = [p.strip() for p in author_field.split(" and ")]
    if any(p.lower() == "others" for p in parts):
        parts = [p for p in parts if p.lower() != "others"]
        if len(parts) >= 1:
            return ", ".join(parts[:6]) + ", et al."
    if len(parts) > 6:
        parts = parts[:6] + ["et al."]
    return ", ".join(parts)


def format_reference(n: int, entry: dict[str, str]) -> str:
    authors = format_authors(entry.get("author"))
    title = entry.get("title", "").rstrip(".")
    journal = entry.get("journal", "")
    year = entry.get("year", "")
    volume = entry.get("volume", "")
    pages = entry.get("pages", "")
    doi = entry.get("doi", "")
    bits = []
    if authors:
        bits.append(authors + ".")
    if title:
        bits.append(title + ".")
    tail = journal
    if year:
        tail += f" {year}"
    if volume:
        tail += f";{volume}"
    if pages:
        tail += f":{pages}"
    if tail.strip():
        bits.append(tail + ".")
    if doi:
        bits.append(f"doi:{doi}.")
    return f"{n}. " + " ".join(bits)


def build_submission_md() -> tuple[str, list[str]]:
    raw = MANUSCRIPT.read_text(encoding="utf-8-sig")
    body = clean_md_body(raw)
    body = insert_figure_callouts(body)
    order = extract_citation_order(body)
    body = replace_citations(body, order)
    bib = parse_bibtex(BIB)
    references = [format_reference(i + 1, bib[k]) for i, k in enumerate(order)]

    abstract_text = "\n".join(f"**{k}:** {v}" for k, v in ABSTRACT.items())
    highlights = "\n".join(f"- {h}" for h in HIGHLIGHTS)
    keywords = "; ".join(KEYWORDS)
    abbreviations = "\n".join(f"{abbr}: {definition}" for abbr, definition in ABBREVIATIONS)
    declarations = []
    for heading, text in DECLARATIONS.items():
        declarations.extend([f"### {heading}", "", text, ""])
    front = [
        f"# {TITLE}",
        "",
        AUTHORS,
        "",
        *AFFILIATIONS,
        "",
        "## Highlights",
        "",
        highlights,
        "",
        "## Abstract",
        "",
        abstract_text,
        "",
        f"**Keywords:** {keywords}",
        "",
        "## Abbreviations",
        "",
        abbreviations,
        "",
    ]
    declarations_section = ["", "## Declarations", "", *declarations]
    refs = ["", "## References", "", *references, ""]
    legends = ["", "## Figure Legends", ""]
    for num in sorted(FIGURES):
        legends.append(f"**Figure {num}. {FIGURES[num]['title']}**")
        legends.append(FIGURES[num]["legend"])
        legends.append("")
    full = "\n".join(front) + "\n" + body + "\n" + "\n".join(declarations_section + refs + legends)
    OUT_MD.write_text(full, encoding="utf-8")
    return full, references


def set_cell_text(paragraph, text, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(0)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_page_break(doc: Document):
    doc.add_page_break()


def add_page_numbers(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def build_docx(submission_md: str, references: list[str]):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    add_page_numbers(sec)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)

    p = add_paragraph(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16)
    p.paragraph_format.line_spacing = 2

    p = add_paragraph(doc, AUTHORS)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for aff in AFFILIATIONS:
        add_paragraph(doc, aff)

    add_heading(doc, "Highlights", 1)
    for h in HIGHLIGHTS:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run("• " + h)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    add_heading(doc, "Abstract", 1)
    for k, v in ABSTRACT.items():
        p = add_paragraph(doc)
        rb = p.add_run(k + ": ")
        rb.bold = True
        rb.font.name = "Times New Roman"
        rb._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rb.font.size = Pt(12)
        rn = p.add_run(v)
        rn.font.name = "Times New Roman"
        rn._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rn.font.size = Pt(12)
    add_paragraph(doc, "Keywords: " + "; ".join(KEYWORDS), bold=False)
    add_page_break(doc)

    # Body from submission MD after front matter and before references.
    body = submission_md.split("**Keywords:**", 1)[1]
    body = body.split("## References", 1)[0]
    body = re.sub(r"^.*?\n\n", "", body, count=1, flags=re.S).strip()
    body = re.sub(r"^\s*# .+?\r?\n+", "", body).replace("\ufeff", "").strip()
    for block in body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("## "):
            add_heading(doc, block[3:].strip(), 1)
        elif block.startswith("### "):
            add_heading(doc, block[4:].strip(), 2)
        elif block.startswith("[Insert Figure"):
            add_paragraph(doc, block, italic=True)
        else:
            add_paragraph(doc, block)

    add_page_break(doc)
    add_heading(doc, "References", 1)
    for ref in references:
        p = add_paragraph(doc, ref)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.25)

    add_page_break(doc)
    add_heading(doc, "Figure Legends", 1)
    for num in sorted(FIGURES):
        add_paragraph(doc, f"Figure {num}. {FIGURES[num]['title']}", bold=True)
        add_paragraph(doc, FIGURES[num]["legend"])

    add_page_break(doc)
    add_heading(doc, "Figures", 1)
    for num in sorted(FIGURES):
        add_paragraph(doc, f"Figure {num}", bold=True)
        if FIGURES[num]["path"].exists():
            doc.add_picture(str(FIGURES[num]["path"]), width=Inches(6.4))
        if num != max(FIGURES):
            add_page_break(doc)

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    submission_md, refs = build_submission_md()
    build_docx(submission_md, refs)
    print(OUT_MD)
    print(OUT_DOCX)

