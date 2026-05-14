from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(r"D:\codex\GenomicSEM\metabolic")
MANUSCRIPT_DIR = ROOT / "manuscript"
SUPP_METHODS_SRC = MANUSCRIPT_DIR / "metabolic_factor_ndd_supp_methods_draft_v3_submission_table_aligned.md"
SUPP_INFO_SRC = MANUSCRIPT_DIR / "metabolic_factor_ndd_supplementary_information_results_v2_clean.md"
SUPP_METHODS_MD = MANUSCRIPT_DIR / "metabolic_factor_ndd_supp_methods_submission_final.md"
SUPP_INFO_MD = MANUSCRIPT_DIR / "metabolic_factor_ndd_supplementary_information_submission_final.md"
SUPP_METHODS_DOCX = MANUSCRIPT_DIR / "metabolic_factor_ndd_supp_methods_submission_final.docx"
SUPP_INFO_DOCX = MANUSCRIPT_DIR / "metabolic_factor_ndd_supplementary_information_submission_final.docx"
SUPP_TABLES_SRC = (
    ROOT
    / "postgwas_ad_pdlbd"
    / "results"
    / "22_supplement_tables_lipid8_F2_AD"
    / "metabolic_factor_triplet_supplementary_tables_v13_S24b_submission.xlsx"
)
SUPP_TABLES_OUT = MANUSCRIPT_DIR / "metabolic_factor_ndd_supplementary_tables_submission_final.xlsx"


SUPP_REFS = [
    "Tambets, Reigo, et al. Mapping the genetic regulation of the blood metabolome at biobank scale. medRxiv 2024. doi:10.1101/2024.10.15.24315557.",
    "Bycroft, Clare, Freeman, Colin, Petkova, Desislava, et al. The UK Biobank resource with deep phenotyping and genomic data. Nature 2018;562:203-209. doi:10.1038/s41586-018-0579-z.",
    "Karjalainen, Juha, et al. Genome-wide characterization of circulating metabolic biomarkers. Nature 2024;628:130-138. doi:10.1038/s41586-024-07148-y.",
    "Bellenguez, Celine, et al. New insights into the genetic etiology of Alzheimer's disease and related dementias. Nature Genetics 2022;54:412-436. doi:10.1038/s41588-022-01024-z.",
    "Kim, Joo-Yeon, et al. Multi-ancestry genome-wide association meta-analysis of Parkinson's disease. Nature Genetics 2024. doi:10.1038/s41588-024-01669-0.",
    "Chia, Ruth, et al. Genome sequencing analysis identifies new loci associated with Lewy body dementia and provides insights into its genetic architecture. Nature Genetics 2021;53:294-303. doi:10.1038/s41588-021-00785-3.",
    "Bulik-Sullivan, Brendan K., Loh, Po-Ru, Finucane, Hilary K., Ripke, Stephan, Yang, Jian, Patterson, Nick, et al. LD Score regression distinguishes confounding from polygenicity in genome-wide association studies. Nature Genetics 2015;47:291-295. doi:10.1038/ng.3211.",
    "Grotzinger, Andrew D., Rhemtulla, Mijke, de Vlaming, Ronald, Ritchie, Stuart J., Mallard, Travis T., Hill, W. David, et al. Genomic structural equation modelling provides insights into the multivariate genetic architecture of complex traits. Nature Human Behaviour 2019;3:513-525. doi:10.1038/s41562-019-0566-x.",
    "Benjamini, Yoav, Hochberg, Yosef. Controlling the false discovery rate: a practical and powerful approach to multiple testing. Journal of the Royal Statistical Society: Series B 1995;57:289-300. doi:10.1111/j.2517-6161.1995.tb02031.x.",
    "Frei, Oleksandr, Holland, Dominic, Smeland, Olav B., et al. Bivariate causal mixture model quantifies polygenic overlap between complex traits beyond genetic correlation. Nature Communications 2019;10:2417. doi:10.1038/s41467-019-10310-0.",
    "1000 Genomes Project Consortium. A global reference for human genetic variation. Nature 2015;526:68-74. doi:10.1038/nature15393.",
    "Andreassen, Ole A., Thompson, Wesley K., Schork, Andrew J., et al. Improved detection of common variants associated with schizophrenia and bipolar disorder using pleiotropy-informed conditional false discovery rate. PLoS Genetics 2013;9:e1003455. doi:10.1371/journal.pgen.1003455.",
    "Giambartolomei, Claudia, Vukcevic, Damjan, Schadt, Eric E., Franke, Lude, Hingorani, Aroon D., Wallace, Chris, et al. Bayesian test for colocalisation between pairs of genetic association studies using summary statistics. PLOS Genetics 2014;10:e1004383. doi:10.1371/journal.pgen.1004383.",
    "Wallace, Chris. A more accurate method for colocalisation analysis allowing for multiple causal variants. PLoS Genetics 2021;17:e1009440. doi:10.1371/journal.pgen.1009440.",
    "Watanabe, Kyoko, Taskesen, Erdogan, van Bochoven, Arjen, Posthuma, Danielle. Functional mapping and annotation of genetic associations with FUMA. Nature Communications 2017;8:1826. doi:10.1038/s41467-017-01261-5.",
    "Zhao, Siming, Crouse, Wesley, Qian, Sheng, Luo, Kaixuan, Stephens, Matthew, He, Xin. Adjusting for genetic confounders in transcriptome-wide association studies improves discovery of risk genes of complex traits. Nature Genetics 2024;56:336-347. doi:10.1038/s41588-023-01648-9.",
    "Zhu, Zhihong, Zhang, Futao, Hu, Han, Bakshi, Andrew, Robinson, Matthew R., Powell, Joseph E., et al. Integration of summary data from GWAS and eQTL studies predicts complex trait gene targets. Nature Genetics 2016;48:481-487. doi:10.1038/ng.3538.",
    "GTEx Consortium. The GTEx Consortium atlas of genetic regulatory effects across human tissues. Science 2020;369:1318-1330. doi:10.1126/science.aaz1776.",
    "Bryois, Julien, et al. Cell-type-specific cis-eQTLs in eight human brain cell types identify novel risk genes for psychiatric and neurological disorders. Nature Neuroscience 2022;25:1104-1112. doi:10.1038/s41593-022-01128-z.",
    "Ma, Yang, et al. A single-cell pathway-based polygenic risk score method for integrating single-cell RNA-seq and GWAS data. Cell Genomics 2023. doi:10.1016/j.xgen.2023.100383.",
    "Osorio, Daniel, Zhong, Yan, Li, Gang, Cai, James J. scTenifoldKnk: An efficient virtual knockout tool for gene function predictions via single-cell gene regulatory network perturbation. Patterns 2022. doi:10.1016/j.patter.2022.100434.",
    "The Open Targets Platform: supporting systematic drug-target identification and prioritisation. Nucleic Acids Research 2025. doi:10.1093/nar/gkae1026.",
    "Freshour, Sharon L., Kiwala, Soma, Cotto, Kelsy C., et al. Integration of the Drug-Gene Interaction Database (DGIdb 4.0) with open crowdsource efforts. Nucleic Acids Research 2021;49:D1144-D1151. doi:10.1093/nar/gkaa1084.",
    "Uhlen, Mathias, Fagerberg, Linn, Hallstrom, Bjorn M., et al. Tissue-based map of the human proteome. Science 2015;347:1260419. doi:10.1126/science.1260419.",
    "Jumper, John, Evans, Richard, Pritzel, Alexander, et al. Highly accurate protein structure prediction with AlphaFold. Nature 2021;596:583-589. doi:10.1038/s41586-021-03819-2.",
    "Tunyasuvunakool, Kathryn, Adler, Jonas, Wu, Zachary, et al. Highly accurate protein structure prediction for the human proteome. Nature 2021;596:590-596. doi:10.1038/s41586-021-03828-1.",
    "Pichet Binette, Alexa, et al. Proteomic changes in Alzheimer's disease associated with progressive amyloid-beta plaque and tau tangle pathologies. Nature Neuroscience 2024;27:1880-1891. doi:10.1038/s41593-024-01737-w.",
    "van den Hurk, Mark, Lau, Shong, Marchetto, Maria C., Mertens, Jerome, Stern, Shani, Corti, Olga, et al. Druggable transcriptomic pathways revealed in Parkinson's patient-derived midbrain neurons. npj Parkinson's Disease 2022;8:134. doi:10.1038/s41531-022-00400-0.",
]


def normalize_text(s: str) -> str:
    replacements = {
        "×": " x ",
        "\u8133": " x ",
        "Ã—": " x ",
        "¡Á": " x ",
        "Î²": "beta",
        "β": "beta",
        "π": "pi",
        "ρ": "rho",
        "–": "-",
        "—": "-",
        "−": "-",
        "≥": ">=",
        "≤": "<=",
        "α": "alpha",
    }
    for old, new in replacements.items():
        s = s.replace(old, new)
    s = re.sub(r"\s+x\s+", " x ", s)
    return s


def patch_supp_methods() -> str:
    text = SUPP_METHODS_SRC.read_text(encoding="utf-8")
    text = text.replace(
        "The primary metabolic GWAS resource was the Estonian Biobank-UK Biobank European meta-analysis release used locally as GCST90451xxx summary statistics.",
        "The primary metabolic GWAS resource was the Estonian Biobank-UK Biobank European meta-analysis release used locally as GCST90451xxx summary statistics [1,2].",
    )
    text = text.replace(
        "The external metabolite validation layer used the independent NMR metabolite GWAS by Karjalainen et al. (Nature, 2024), not the same discovery release.",
        "The external metabolite validation layer used the independent NMR metabolite GWAS by Karjalainen et al. (Nature, 2024), not the same discovery release [3].",
    )
    text = text.replace(
        "The AD file corresponded to the European Alzheimer & Dementia Biobank stage-I GWAS by Bellenguez et al., including 39,106 clinically diagnosed AD cases, 46,828 proxy-AD cases, and 401,577 controls.",
        "The AD file corresponded to the European Alzheimer & Dementia Biobank stage-I GWAS by Bellenguez et al., including 39,106 clinically diagnosed AD cases, 46,828 proxy-AD cases, and 401,577 controls [4].",
    )
    text = text.replace(
        "The PD file corresponded to the local European-ancestry comparison summary statistics from the multi-ancestry PD GWAS meta-analysis by Kim et al., with 63,555 cases, 17,700 proxy cases, and 1,746,386 controls.",
        "The PD file corresponded to the local European-ancestry comparison summary statistics from the multi-ancestry PD GWAS meta-analysis by Kim et al., with 63,555 cases, 17,700 proxy cases, and 1,746,386 controls [5].",
    )
    text = text.replace(
        "The LBD file corresponded to the genome-sequencing association study by Chia et al., with 2,591 cases and 4,027 controls.",
        "The LBD file corresponded to the genome-sequencing association study by Chia et al., with 2,591 cases and 4,027 controls [6].",
    )
    text = text.replace(
        "The main modeling universe excluded proportion-only traits and retained traits with robust SNP heritability signal by univariate LDSC.",
        "The main modeling universe excluded proportion-only traits and retained traits with robust SNP heritability signal by univariate LDSC [7].",
    )
    text = text.replace(
        "Multivariable LDSC was used to estimate the genetic covariance matrix among retained indicators.",
        "Multivariable LDSC was used to estimate the genetic covariance matrix among retained indicators [7,8].",
    )
    text = text.replace(
        "Final model evaluation used Genomic SEM `usermodel()` on retained indicators only, with formal fit statistics reported in Supplementary Table S6 and standardized loadings reported in Supplementary Table S7.",
        "Final model evaluation used Genomic SEM `usermodel()` on retained indicators only [8], with formal fit statistics reported in Supplementary Table S6 and standardized loadings reported in Supplementary Table S7.",
    )
    text = text.replace(
        "Factor GWAS were generated with the local Genomic SEM workflow using `GenomicSEM::sumstats()` and `GenomicSEM::userGWAS()`.",
        "Factor GWAS were generated with the local Genomic SEM workflow using `GenomicSEM::sumstats()` and `GenomicSEM::userGWAS()` [8].",
    )
    text = text.replace(
        "using the 1000 Genomes European reference with `P < 5e-8`, `r2 = 0.1`, and `kb = 1000`",
        "using the 1000 Genomes European reference [11] with `P < 5e-8`, `r2 = 0.1`, and `kb = 1000`",
    )
    text = text.replace(
        "This interpretation is consistent with the role of Q_SNP in Genomic SEM as a test of SNP-level heterogeneity relative to the specified factor structure.",
        "This interpretation is consistent with the role of Q_SNP in Genomic SEM as a test of SNP-level heterogeneity relative to the specified factor structure [8].",
    )
    text = text.replace(
        "Each exported factor GWAS was re-evaluated with univariate LDSC.",
        "Each exported factor GWAS was re-evaluated with univariate LDSC [7].",
    )
    text = text.replace(
        "Bivariate LDSC was run between factor GWAS and the retained indicator/support traits, and the results were summarized in compact and full-pair tables",
        "Bivariate LDSC was run between factor GWAS and the retained indicator/support traits [7], and the results were summarized in compact and full-pair tables",
    )
    text = text.replace(
        "External validation used the independent Karjalainen et al. NMR metabolite GWAS atlas.",
        "External validation used the independent Karjalainen et al. NMR metabolite GWAS atlas [3].",
    )
    text = text.replace(
        "All six factor GWAS were compared with AD, PD, and LBD in a complete 18-pair LDSC screen",
        "All six factor GWAS were compared with AD, PD, and LBD in a complete 18-pair LDSC screen using bivariate LDSC [7]",
    )
    text = text.replace(
        "Bivariate MiXeR was applied to the three focal factor-disease pairs to quantify polygenic overlap beyond LDSC genetic correlation.",
        "Bivariate MiXeR was applied to the three focal factor-disease pairs to quantify polygenic overlap beyond LDSC genetic correlation [10].",
    )
    text = text.replace(
        "the 1000 Genomes European Phase 3 PLINK reference",
        "the 1000 Genomes European Phase 3 PLINK reference [11]",
    )
    text = text.replace(
        "pleiotropy-informed conjunctional false discovery rate analyses were used",
        "pleiotropy-informed conjunctional false discovery rate analyses were used [12]",
    )
    text = text.replace(
        "and retained loci passing `conjFDR < 0.05`.",
        "and retained loci passing `conjFDR < 0.05`; false-discovery-rate terminology follows the Benjamini-Hochberg framework [9].",
    )
    text = text.replace(
        "Regional colocalization was performed with `coloc.abf`.",
        "Regional colocalization was performed with `coloc.abf` [13].",
    )
    text = text.replace(
        "PWCoCo was used as a complementary conditional framework",
        "PWCoCo was used as a complementary conditional framework informed by conditional colocalization principles [14]",
    )
    text = text.replace(
        "FUMA was used for positional and regulatory annotation of prioritized loci",
        "FUMA was used for positional and regulatory annotation of prioritized loci [15]",
    )
    text = text.replace(
        "Single-trait cTWAS was performed separately for factor GWAS and disease GWAS",
        "Single-trait cTWAS was performed separately for factor GWAS and disease GWAS [16]",
    )
    text = text.replace(
        "Summary-data-based Mendelian randomization was used as a direction- and context-aware expression-prioritization layer.",
        "Summary-data-based Mendelian randomization was used as a direction- and context-aware expression-prioritization layer [17].",
    )
    text = text.replace(
        "GTEx v8 brain tissues, and Bryois brain cell-type eQTL panels.",
        "GTEx v8 brain tissues [18], and Bryois brain cell-type eQTL panels [19].",
    )
    text = text.replace(
        "Local summarization scripts applied false-discovery-rate correction to SMR P values",
        "Local summarization scripts applied Benjamini-Hochberg false-discovery-rate correction [9] to SMR P values",
    )
    text = text.replace(
        "Disease-contextual cell mapping was performed with scPagwas",
        "Disease-contextual cell mapping was performed with scPagwas [20]",
    )
    text = text.replace(
        "Production scPagwas runs used `singlecell = TRUE`, `celltype = TRUE`, and `iters_singlecell = 100`.",
        "Production scPagwas runs used `singlecell = TRUE`, `celltype = TRUE`, and `iters_singlecell = 100` [20].",
    )
    text = text.replace(
        "and adjusted within analysis using the Benjamini-Hochberg procedure.",
        "and adjusted within analysis using the Benjamini-Hochberg procedure [9].",
    )
    text = text.replace(
        "scTenifoldKnk analyses were restricted",
        "scTenifoldKnk analyses [21] were restricted",
    )
    text = text.replace(
        "The production KNK scripts used `scTenifoldKnk` and `scTenifoldNet`.",
        "The production KNK scripts used `scTenifoldKnk` and `scTenifoldNet` [21].",
    )
    text = text.replace(
        "The annotation combined Open Targets target and disease-knowledge fields, DGIdb drug-gene interaction records, Human Protein Atlas brain-expression evidence, Open Targets interaction/network information, chemical-probe evidence, and AlphaFold structure availability.",
        "The annotation combined Open Targets target and disease-knowledge fields [22], DGIdb drug-gene interaction records [23], Human Protein Atlas brain-expression evidence [24], Open Targets interaction/network information [22], chemical-probe evidence, and AlphaFold structure availability [25,26].",
    )
    text = text.replace(
        "For AD, we used processed CSF proteomic supplementary tables from Binette et al. (Nature Neuroscience, 2024)",
        "For AD, we used processed CSF proteomic supplementary tables from Binette et al. (Nature Neuroscience, 2024) [27]",
    )
    text = text.replace(
        "For PD, we used the processed transcriptomic meta-signature from van den Hurk et al. (npj Parkinson's Disease, 2022)",
        "For PD, we used the processed transcriptomic meta-signature from van den Hurk et al. (npj Parkinson's Disease, 2022) [28]",
    )
    text = re.sub(r"\n## 21\. Citation notes for reference formatting\n\n.*\Z", "", text, flags=re.S)
    text += "\n\n## References\n\n"
    text += "\n".join(f"{i}. {ref}" for i, ref in enumerate(SUPP_REFS, start=1))
    return text


def patch_supp_info() -> str:
    text = SUPP_INFO_SRC.read_text(encoding="utf-8")
    text = text.replace(
        "112 non-proportion traits from the primary EstBB-UKBB `meta_EUR` NMR metabolite GWAS",
        "112 non-proportion traits from the primary EstBB-UKBB `meta_EUR` NMR metabolite GWAS [1,2]",
    )
    text = text.replace("univariate LDSC entry rule", "univariate LDSC entry rule [4]")
    text = text.replace("Genomic SEM Q_SNP", "Genomic SEM Q_SNP [5]")
    text = text.replace(
        "External validation used the Karjalainen et al. NMR metabolite GWAS resource",
        "External validation used the Karjalainen et al. NMR metabolite GWAS resource [3]",
    )
    text = text.replace(
        "In the PD transcriptomic meta-signature",
        "In the PD transcriptomic meta-signature [7]",
    )
    text = text.replace(
        "The public AD CSF proteomic resource",
        "The public AD CSF proteomic resource [6]",
    )
    text += "\n\n## References\n\n"
    used = [SUPP_REFS[i - 1] for i in [1, 2, 3, 7, 8, 27, 28]]
    text += "\n".join(f"{i}. {ref}" for i, ref in enumerate(used, start=1))
    return text


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field in Word to refresh page numbers."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def collect_headings(md_text: str) -> list[tuple[int, str]]:
    headings = []
    for line in md_text.splitlines():
        if line.startswith("### "):
            headings.append((3, line[4:].strip()))
        elif line.startswith("## "):
            headings.append((2, line[3:].strip()))
        elif line.startswith("# "):
            headings.append((1, line[2:].strip()))
    return headings


def set_update_fields(docx_path: Path) -> None:
    tmp = Path(tempfile.mkdtemp(prefix="supp_docx_settings_"))
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            z.extractall(tmp)
        settings_path = tmp / "word" / "settings.xml"
        if settings_path.exists():
            settings = settings_path.read_text(encoding="utf-8")
        else:
            settings_path.parent.mkdir(parents=True, exist_ok=True)
            settings = '<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:settings>'
        if "w:updateFields" not in settings:
            settings = settings.replace(
                "</w:settings>",
                '<w:updateFields w:val="true"/></w:settings>',
            )
            settings_path.write_text(settings, encoding="utf-8")
        with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for p in tmp.rglob("*"):
                if p.is_file():
                    z.write(p, p.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def add_md_runs(paragraph, text: str) -> None:
    text = normalize_text(text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def md_to_docx(md_text: str, out_docx: Path, title: str) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Arial"
    doc.add_paragraph("Table of Contents", style="Heading 1")
    for level, heading in collect_headings(md_text):
        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.left_indent = Inches(0.18 * max(level - 1, 0))
        toc_run = toc_p.add_run(heading)
        toc_run.font.name = "Times New Roman"
        toc_run.font.size = Pt(10.5)
        if level <= 2:
            toc_run.bold = True
    doc.add_page_break()
    for block in re.split(r"\n\s*\n", md_text.strip()):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            doc.add_paragraph(block[2:].strip(), style="Heading 1")
        elif block.startswith("## "):
            doc.add_paragraph(block[3:].strip(), style="Heading 2")
        elif block.startswith("### "):
            doc.add_paragraph(block[4:].strip(), style="Heading 3")
        else:
            for line in block.splitlines():
                line = line.strip()
                if not line:
                    continue
                para = doc.add_paragraph()
                if re.match(r"^\d+\. ", line):
                    para.paragraph_format.left_indent = Inches(0.22)
                    para.paragraph_format.first_line_indent = Inches(-0.22)
                add_md_runs(para, line)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.line_spacing = 1.08
    doc.save(out_docx)
    set_update_fields(out_docx)


def clean_supplementary_tables() -> None:
    wb = load_workbook(SUPP_TABLES_SRC)
    index = wb["Supplementary_Table_Index"]
    titles = {}
    for row in index.iter_rows(min_row=2, values_only=True):
        if row[0]:
            titles[str(row[0])] = (str(row[1] or ""), str(row[2] or ""))
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    title_fill = PatternFill("solid", fgColor="1F4E79")
    note_fill = PatternFill("solid", fgColor="EAF3F8")
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = normalize_text(cell.value)
        ws.freeze_panes = "A4"
        ws.sheet_view.showGridLines = False
        max_col = ws.max_column
        if ws.title == "Supplementary_Table_Index":
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = title_fill
                cell.alignment = Alignment(wrap_text=True, vertical="center")
            ws.column_dimensions["A"].width = 32
            ws.column_dimensions["B"].width = 48
            ws.column_dimensions["C"].width = 88
            continue
        table_id = ws.title.split("_", 1)[0].replace("S", "S")
        title, desc = titles.get(ws.title, (ws.title, ""))
        ws.cell(1, 1).value = f"Supplementary Table {table_id}. {normalize_text(title)}"
        ws.cell(2, 1).value = f"Note: {normalize_text(desc)}"
        for row_idx, fill, color in [(1, title_fill, "FFFFFF"), (2, note_fill, "000000")]:
            ws.cell(row_idx, 1).font = Font(bold=(row_idx == 1), color=color)
            ws.cell(row_idx, 1).fill = fill
            ws.cell(row_idx, 1).alignment = Alignment(wrap_text=True, vertical="center")
        if max_col > 1:
            try:
                ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
                ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_col)
            except ValueError:
                pass
        ws.row_dimensions[1].height = 24
        ws.row_dimensions[2].height = 36
        # Style the first row after the title block that appears to contain column names.
        header_row = None
        for r in range(3, min(ws.max_row, 10) + 1):
            values = [ws.cell(r, c).value for c in range(1, max_col + 1)]
            filled = sum(v is not None for v in values)
            if filled >= max(2, min(4, max_col)):
                header_row = r
                break
        if header_row:
            for cell in ws[header_row]:
                cell.font = Font(bold=True)
                cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=True, vertical="center")
        for col in range(1, max_col + 1):
            letter = get_column_letter(col)
            max_len = 0
            for row in range(1, min(ws.max_row, 80) + 1):
                val = ws.cell(row, col).value
                if val is not None:
                    max_len = max(max_len, len(str(val)))
            ws.column_dimensions[letter].width = max(10, min(42, max_len + 2))
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
    wb.save(SUPP_TABLES_OUT)


def main() -> None:
    supp_methods = patch_supp_methods()
    SUPP_METHODS_MD.write_text(supp_methods, encoding="utf-8")
    supp_info = patch_supp_info()
    SUPP_INFO_MD.write_text(supp_info, encoding="utf-8")
    md_to_docx(supp_methods, SUPP_METHODS_DOCX, "Supplementary Methods")
    md_to_docx(supp_info, SUPP_INFO_DOCX, "Supplementary Information")
    clean_supplementary_tables()
    print(SUPP_METHODS_MD)
    print(SUPP_METHODS_DOCX)
    print(SUPP_INFO_MD)
    print(SUPP_INFO_DOCX)
    print(SUPP_TABLES_OUT)


if __name__ == "__main__":
    main()
